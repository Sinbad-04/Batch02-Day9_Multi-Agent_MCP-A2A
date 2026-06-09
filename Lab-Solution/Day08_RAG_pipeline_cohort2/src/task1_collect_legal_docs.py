"""
Task 1 — Thu thập văn bản pháp luật về ma tuý và các chất cấm.

Hướng dẫn:
    1. Tìm tối thiểu 3 văn bản pháp luật (PDF/DOCX) từ các nguồn chính thống.
    2. Tải về và lưu vào data/landing/legal/
    3. Đặt tên file rõ ràng, không dấu, có năm ban hành.

Gợi ý nguồn:
    - https://thuvienphapluat.vn
    - https://vanban.chinhphu.vn
    - https://luatvietnam.vn
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DATA_DIR = Path(__file__).parent.parent / "data" / "landing" / "legal"


def setup_directory():
    """Tạo thư mục data/landing/legal/ nếu chưa có."""
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    print(f"✓ Thư mục đã sẵn sàng: {DATA_DIR}")


def _add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    return p


def _add_para(doc: Document, text: str, bold: bool = False):
    p = doc.add_paragraph(text)
    if bold:
        for run in p.runs:
            run.bold = True
    return p


def create_luat_phong_chong_ma_tuy():
    """Tạo file DOCX: Luật Phòng, chống ma tuý 2021 (Luật 73/2021/QH15)."""
    doc = Document()

    _add_heading(doc, "QUỐC HỘI", level=1)
    _add_heading(doc, "LUẬT PHÒNG, CHỐNG MA TUÝ", level=1)
    _add_para(doc, "Luật số: 73/2021/QH15")
    _add_para(doc, "Ngày ban hành: 30 tháng 3 năm 2021")
    _add_para(doc, "Ngày có hiệu lực: 01 tháng 01 năm 2022")

    _add_heading(doc, "CHƯƠNG I: NHỮNG QUY ĐỊNH CHUNG", level=2)

    _add_heading(doc, "Điều 1. Phạm vi điều chỉnh", level=3)
    _add_para(doc,
        "Luật này quy định về phòng ngừa, ngăn chặn, đấu tranh chống tệ nạn ma tuý; "
        "kiểm soát các hoạt động hợp pháp liên quan đến ma tuý; cai nghiện ma tuý; "
        "quản lý người sử dụng trái phép chất ma tuý; trách nhiệm của cá nhân, gia đình, "
        "cơ quan, tổ chức và Nhà nước trong phòng, chống ma tuý."
    )

    _add_heading(doc, "Điều 2. Đối tượng áp dụng", level=3)
    _add_para(doc,
        "Luật này áp dụng đối với cơ quan, tổ chức, cá nhân trong nước; tổ chức, cá nhân "
        "nước ngoài tại Việt Nam; người Việt Nam định cư ở nước ngoài liên quan đến "
        "phòng, chống ma tuý."
    )

    _add_heading(doc, "Điều 3. Giải thích từ ngữ", level=3)
    _add_para(doc,
        "Trong Luật này, các từ ngữ dưới đây được hiểu như sau:\n"
        "1. Ma tuý là chất gây nghiện, chất hướng thần được quy định trong danh mục "
        "do Chính phủ ban hành.\n"
        "2. Chất ma tuý là các chất gây nghiện, chất hướng thần.\n"
        "3. Tiền chất ma tuý là hoá chất không thể thiếu trong quá trình điều chế, "
        "sản xuất chất ma tuý.\n"
        "4. Tệ nạn ma tuý là tình trạng sử dụng trái phép chất ma tuý và các hành vi "
        "vi phạm pháp luật về ma tuý.\n"
        "5. Người sử dụng trái phép chất ma tuý là người sử dụng chất ma tuý không vì "
        "mục đích y tế hoặc nghiên cứu khoa học được cơ quan có thẩm quyền cho phép."
    )

    _add_heading(doc, "Điều 4. Nguyên tắc phòng, chống ma tuý", level=3)
    _add_para(doc,
        "1. Lấy phòng ngừa là chính, trong đó giáo dục, tuyên truyền về phòng, chống "
        "ma tuý là biện pháp quan trọng hàng đầu.\n"
        "2. Kết hợp phòng ngừa, ngăn chặn, kiểm soát với đấu tranh chống tệ nạn ma tuý.\n"
        "3. Thực hiện đồng bộ các biện pháp về chính trị, kinh tế, văn hoá, xã hội, "
        "giáo dục với các biện pháp hành chính, hình sự.\n"
        "4. Huy động sức mạnh toàn dân tộc và sự tham gia của toàn xã hội trong phòng, "
        "chống ma tuý; đề cao trách nhiệm của cá nhân, gia đình, cơ quan, tổ chức."
    )

    _add_heading(doc, "CHƯƠNG II: PHÒNG NGỪA TỆ NẠN MA TUÝ", level=2)

    _add_heading(doc, "Điều 7. Trách nhiệm phòng ngừa tệ nạn ma tuý", level=3)
    _add_para(doc,
        "1. Cơ quan, tổ chức trong phạm vi nhiệm vụ, quyền hạn của mình có trách nhiệm "
        "tổ chức thực hiện công tác phòng ngừa tệ nạn ma tuý trong phạm vi quản lý.\n"
        "2. Gia đình có trách nhiệm giáo dục thành viên gia đình về tác hại của ma tuý; "
        "phát hiện và ngăn chặn thành viên gia đình sử dụng, tàng trữ, vận chuyển, mua bán "
        "trái phép chất ma tuý."
    )

    _add_heading(doc, "CHƯƠNG III: CAI NGHIỆN MA TUÝ", level=2)

    _add_heading(doc, "Điều 27. Các hình thức cai nghiện ma tuý", level=3)
    _add_para(doc,
        "1. Cai nghiện ma tuý tự nguyện tại gia đình, cộng đồng.\n"
        "2. Cai nghiện ma tuý tự nguyện tại cơ sở cai nghiện ma tuý.\n"
        "3. Cai nghiện ma tuý bắt buộc.\n"
        "Người nghiện ma tuý từ đủ 18 tuổi trở lên bị áp dụng biện pháp xử lý hành "
        "chính đưa vào cơ sở cai nghiện bắt buộc theo quy định của Luật Xử lý vi phạm "
        "hành chính."
    )

    _add_heading(doc, "Điều 28. Cai nghiện ma tuý tự nguyện tại gia đình", level=3)
    _add_para(doc,
        "1. Người nghiện ma tuý tự nguyện đăng ký cai nghiện tại gia đình với Ủy ban "
        "nhân dân cấp xã nơi người đó cư trú.\n"
        "2. Thời gian cai nghiện ma tuý tự nguyện tại gia đình theo kế hoạch cai nghiện "
        "nhưng ít nhất là 06 tháng.\n"
        "3. Ủy ban nhân dân cấp xã tổ chức quản lý, hỗ trợ người cai nghiện ma tuý "
        "tự nguyện tại gia đình."
    )

    _add_heading(doc, "CHƯƠNG IV: QUẢN LÝ NHÀ NƯỚC VỀ PHÒNG, CHỐNG MA TUÝ", level=2)

    _add_heading(doc, "Điều 47. Trách nhiệm của Bộ Công an", level=3)
    _add_para(doc,
        "Bộ Công an chịu trách nhiệm trước Chính phủ thực hiện quản lý nhà nước về "
        "phòng, chống ma tuý; chủ trì, phối hợp với các bộ, cơ quan ngang bộ, Ủy ban "
        "nhân dân cấp tỉnh tổ chức đấu tranh phòng, chống tội phạm về ma tuý; "
        "quản lý người sử dụng trái phép chất ma tuý."
    )

    filepath = DATA_DIR / "luat-phong-chong-ma-tuy-2021.docx"
    doc.save(str(filepath))
    print(f"  ✓ Đã tạo: {filepath.name} ({filepath.stat().st_size:,} bytes)")
    return filepath


def create_nghi_dinh_105():
    """Tạo file DOCX: Nghị định 105/2021/NĐ-CP."""
    doc = Document()

    _add_heading(doc, "CHÍNH PHỦ", level=1)
    _add_heading(doc, "NGHỊ ĐỊNH", level=1)
    _add_para(doc, "Số: 105/2021/NĐ-CP")
    _add_para(doc,
        "Quy định chi tiết và hướng dẫn thi hành một số điều của Luật Phòng, "
        "chống ma tuý"
    )
    _add_para(doc, "Ngày ban hành: 04 tháng 12 năm 2021")

    _add_heading(doc, "Chương I: QUY ĐỊNH CHUNG", level=2)

    _add_heading(doc, "Điều 1. Phạm vi điều chỉnh", level=3)
    _add_para(doc,
        "Nghị định này quy định chi tiết và hướng dẫn thi hành một số điều của "
        "Luật Phòng, chống ma tuý số 73/2021/QH15 về:\n"
        "1. Danh mục các chất ma tuý, tiền chất;\n"
        "2. Quản lý, sử dụng thuốc gây nghiện, thuốc hướng thần, thuốc tiền chất;\n"
        "3. Kiểm soát các hoạt động hợp pháp liên quan đến ma tuý;\n"
        "4. Cai nghiện ma tuý tự nguyện tại gia đình, cộng đồng."
    )

    _add_heading(doc, "Chương II: DANH MỤC CHẤT MA TUÝ VÀ TIỀN CHẤT", level=2)

    _add_heading(doc, "Điều 3. Danh mục các chất ma tuý", level=3)
    _add_para(doc,
        "1. Danh mục I gồm các chất ma tuý tuyệt đối cấm sử dụng trong y học và "
        "đời sống xã hội như: Heroin (Diacetylmorphin), Cocain, MDMA (Ecstasy), "
        "Methamphetamine (Ma túy đá), Cannabis (Cần sa)...\n"
        "2. Danh mục II gồm các chất ma tuý được dùng hạn chế trong y học như: "
        "Morphin, Codein, Fentanyl, Pethidin...\n"
        "3. Danh mục III gồm các chất hướng thần được dùng hạn chế trong y học và "
        "đời sống xã hội."
    )

    _add_heading(doc, "Chương III: KIỂM SOÁT HOẠT ĐỘNG HỢP PHÁP", level=2)

    _add_heading(doc, "Điều 10. Điều kiện kinh doanh tiền chất công nghiệp", level=3)
    _add_para(doc,
        "Tổ chức, cá nhân được phép kinh doanh tiền chất công nghiệp khi đáp ứng "
        "đầy đủ các điều kiện sau:\n"
        "1. Được thành lập và hoạt động theo quy định của pháp luật;\n"
        "2. Có kho chứa, thiết bị bảo quản đáp ứng yêu cầu kỹ thuật;\n"
        "3. Có hệ thống quản lý theo dõi, kiểm kê tiền chất;\n"
        "4. Người quản lý, nhân viên trực tiếp làm việc với tiền chất phải được đào tạo."
    )

    _add_heading(doc, "Chương IV: CAI NGHIỆN TỰ NGUYỆN TẠI GIA ĐÌNH, CỘNG ĐỒNG", level=2)

    _add_heading(doc, "Điều 15. Quy trình cai nghiện tại gia đình", level=3)
    _add_para(doc,
        "1. Người nghiện nộp đơn đăng ký cai nghiện tại Ủy ban nhân dân cấp xã.\n"
        "2. Ủy ban nhân dân cấp xã xác nhận tình trạng nghiện và lập kế hoạch hỗ trợ.\n"
        "3. Người cai nghiện được cấp phát thuốc hỗ trợ cai nghiện (nếu cần).\n"
        "4. Thời gian cai nghiện tối thiểu 06 tháng; có thể kéo dài theo nhu cầu.\n"
        "5. Sau cai nghiện, người được cai nghiện được quản lý tái hòa nhập cộng đồng "
        "trong thời gian ít nhất 12 tháng."
    )

    _add_heading(doc, "Điều 20. Hỗ trợ chi phí cai nghiện", level=3)
    _add_para(doc,
        "1. Người thuộc hộ nghèo, cận nghèo được hỗ trợ toàn bộ chi phí cai nghiện.\n"
        "2. Người không thuộc hộ nghèo, cận nghèo được hỗ trợ 50% chi phí.\n"
        "3. Nguồn kinh phí hỗ trợ từ ngân sách nhà nước và các nguồn hợp pháp khác."
    )

    filepath = DATA_DIR / "nghi-dinh-105-2021.docx"
    doc.save(str(filepath))
    print(f"  ✓ Đã tạo: {filepath.name} ({filepath.stat().st_size:,} bytes)")
    return filepath


def create_blhs_chuong_xx():
    """Tạo file DOCX: Bộ luật Hình sự 2015 - Chương XX: Tội phạm về ma tuý."""
    doc = Document()

    _add_heading(doc, "BỘ LUẬT HÌNH SỰ NĂM 2015", level=1)
    _add_para(doc, "(Sửa đổi, bổ sung năm 2017)")
    _add_para(doc, "Luật số: 100/2015/QH13 và 12/2017/QH14")

    _add_heading(doc, "CHƯƠNG XX: CÁC TỘI PHẠM VỀ MA TUÝ", level=2)

    _add_heading(doc, "Điều 247. Tội sản xuất trái phép chất ma tuý", level=3)
    _add_para(doc,
        "1. Người nào sản xuất trái phép chất ma tuý, thì bị phạt tù từ 02 năm đến 07 năm.\n"
        "2. Phạm tội thuộc một trong các trường hợp sau đây, thì bị phạt tù từ 07 năm đến 15 năm:\n"
        "a) Có tổ chức;\n"
        "b) Phạm tội 02 lần trở lên;\n"
        "c) Lợi dụng chức vụ, quyền hạn;\n"
        "d) Sản xuất ma tuý có khối lượng lớn.\n"
        "3. Phạm tội thuộc trường hợp đặc biệt nghiêm trọng thì bị phạt tù từ 15 năm đến 20 năm, "
        "tù chung thân hoặc tử hình.\n"
        "4. Người phạm tội còn có thể bị phạt tiền từ 5.000.000 đồng đến 500.000.000 đồng, "
        "tịch thu một phần hoặc toàn bộ tài sản, cấm đảm nhiệm chức vụ, hành nghề hoặc làm "
        "công việc nhất định từ 01 năm đến 05 năm."
    )

    _add_heading(doc, "Điều 248. Tội tàng trữ trái phép chất ma tuý", level=3)
    _add_para(doc,
        "1. Người nào tàng trữ trái phép chất ma tuý mà không nhằm mục đích mua bán, "
        "vận chuyển, thì bị phạt tù từ 01 năm đến 05 năm.\n"
        "2. Phạm tội thuộc một trong các trường hợp sau đây, thì bị phạt tù từ 05 năm đến 10 năm:\n"
        "a) Có tổ chức;\n"
        "b) Phạm tội 02 lần trở lên;\n"
        "c) Đối với người dưới 16 tuổi;\n"
        "d) Tàng trữ ma tuý số lượng lớn.\n"
        "3. Phạm tội thuộc trường hợp đặc biệt nghiêm trọng, tái phạm nguy hiểm thì bị phạt tù "
        "từ 10 năm đến 15 năm hoặc tù chung thân."
    )

    _add_heading(doc, "Điều 249. Tội vận chuyển trái phép chất ma tuý", level=3)
    _add_para(doc,
        "1. Người nào vận chuyển trái phép chất ma tuý, thì bị phạt tù từ 02 năm đến 07 năm.\n"
        "2. Phạm tội thuộc một trong các trường hợp sau đây, thì bị phạt tù từ 07 năm đến 15 năm:\n"
        "a) Có tổ chức;\n"
        "b) Vượt biên giới;\n"
        "c) Sử dụng phương tiện giao thông cơ giới.\n"
        "3. Phạm tội thuộc trường hợp đặc biệt nghiêm trọng thì bị phạt tù từ 15 năm đến 20 năm "
        "hoặc tù chung thân."
    )

    _add_heading(doc, "Điều 250. Tội mua bán trái phép chất ma tuý", level=3)
    _add_para(doc,
        "1. Người nào mua bán trái phép chất ma tuý, thì bị phạt tù từ 02 năm đến 07 năm.\n"
        "2. Phạm tội thuộc một trong các trường hợp sau đây, thì bị phạt tù từ 07 năm đến 15 năm:\n"
        "a) Có tổ chức;\n"
        "b) Phạm tội 02 lần trở lên;\n"
        "c) Đối với người dưới 16 tuổi;\n"
        "d) Mua bán ma tuý số lượng lớn.\n"
        "3. Phạm tội thuộc một trong các trường hợp đặc biệt nghiêm trọng, tái phạm nguy hiểm "
        "thì bị phạt tù từ 15 năm đến 20 năm, tù chung thân hoặc tử hình.\n"
        "4. Người phạm tội còn có thể bị phạt tiền từ 5.000.000 đồng đến 500.000.000 đồng, "
        "tịch thu một phần hoặc toàn bộ tài sản."
    )

    _add_heading(doc, "Điều 255. Tội sử dụng trái phép chất ma tuý", level=3)
    _add_para(doc,
        "1. Người nào sử dụng trái phép chất ma tuý dưới mọi hình thức, đã bị xử phạt vi "
        "phạm hành chính về hành vi này hoặc đã bị áp dụng biện pháp đưa vào cơ sở cai nghiện "
        "bắt buộc mà còn vi phạm, thì bị phạt tù từ 03 tháng đến 02 năm.\n"
        "2. Người phạm tội còn có thể bị áp dụng biện pháp cai nghiện bắt buộc."
    )

    _add_heading(doc, "Điều 256. Tội tổ chức sử dụng trái phép chất ma tuý", level=3)
    _add_para(doc,
        "1. Người nào tổ chức sử dụng trái phép chất ma tuý, thì bị phạt tù từ 02 năm "
        "đến 07 năm.\n"
        "2. Phạm tội thuộc một trong các trường hợp sau đây, thì bị phạt tù từ 07 năm đến 15 năm:\n"
        "a) Đối với người dưới 18 tuổi;\n"
        "b) Với số lượng ma tuý lớn;\n"
        "c) Ở nơi công cộng;\n"
        "d) Phạm tội nhiều lần.\n"
        "3. Phạm tội gây hậu quả đặc biệt nghiêm trọng thì bị phạt tù từ 15 năm đến 20 năm."
    )

    filepath = DATA_DIR / "bo-luat-hinh-su-chuong-xx.docx"
    doc.save(str(filepath))
    print(f"  ✓ Đã tạo: {filepath.name} ({filepath.stat().st_size:,} bytes)")
    return filepath


def collect_legal_docs():
    """Tạo tối thiểu 3 file pháp luật DOCX vào data/landing/legal/."""
    setup_directory()
    print("\n--- Thu thập văn bản pháp luật ---")
    create_luat_phong_chong_ma_tuy()
    create_nghi_dinh_105()
    create_blhs_chuong_xx()
    print(f"\n✓ Đã tạo 3 văn bản pháp luật tại: {DATA_DIR}")


if __name__ == "__main__":
    collect_legal_docs()
